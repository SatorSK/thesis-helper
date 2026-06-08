"""M7. 내보내기 — 동국대 국제통상학과 양식 .docx 자동 생성.

규정 반영: A4 / 여백 30mm / 머리말·꼬리말 15mm / 본문 11pt·주석 9pt /
줄간격 160% / 표지(제목·학과·학번·지도교수·이름·제출일) / 본문 각주 인용 /
APA 참고문헌 목록 / 파일명 (국제통상학과)학번_이름.docx.

최종 문서에는 M5 '본인 작성분'만 들어간다(초안 제외).
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from core import citation
from core.docx_footnotes import FootnoteManager
from modules.m3_outline import ensure_outline
from modules.m5_drafter import similarity, SIMILARITY_BLOCK


def _setup_format(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(30)
    sec.bottom_margin = Mm(30)
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(30)
    sec.header_distance = Mm(15)
    sec.footer_distance = Mm(15)

    normal = doc.styles["Normal"]
    normal.font.name = "바탕"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.6


def _body_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.6
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def blocked_sections(project) -> list[str]:
    """복붙 차단(유사도≥0.80) 절 목록. 비어 있어야 내보내기 가능."""
    bad = []
    for s in ensure_outline(project):
        para = project["paragraphs"].get(s["id"], {})
        ut = para.get("user_text", "")
        if ut.strip() and similarity(para.get("draft", ""), ut) >= SIMILARITY_BLOCK:
            bad.append(s["title"])
    return bad


def build_docx(project, use_footnotes: bool = True) -> bytes:
    doc = Document()
    _setup_format(doc)
    meta = project["meta"]
    src_by_id = {s["id"]: s for s in project["sources"]}

    # ---- 표지 ----
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(meta.get("title") or project["topic"].get("selected") or "제목 없음")
    r.bold = True
    r.font.size = Pt(20)
    for _ in range(10):
        doc.add_paragraph()
    for label, key in [("학과", "dept"), ("학번", "student_id"),
                       ("지도교수", "advisor"), ("이름", "author"), ("제출일", "date")]:
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.add_run(f"{label}: {meta.get(key,'')}").font.size = Pt(12)
    doc.add_page_break()

    mgr = FootnoteManager(doc) if use_footnotes else None

    # ---- 본문 ----
    for sec in ensure_outline(project):
        h = doc.add_heading(level=min(sec.get("level", 1), 3))
        h.add_run("")  # 스타일 유지
        h.text = sec["title"]
        para_data = project["paragraphs"].get(sec["id"], {})
        body = (para_data.get("user_text") or "").strip()
        if not body:
            _body_para(doc, "(본문 미작성)")
            continue
        for chunk in body.split("\n"):
            if chunk.strip():
                _body_para(doc, chunk.strip())
        # 각주 삽입: 이 절에 연결된 출처
        last_p = doc.paragraphs[-1]
        for cid in para_data.get("cite_ids", []):
            s = src_by_id.get(cid)
            if not s:
                continue
            note = citation.footnote(s)
            if mgr is not None:
                mgr.add(last_p, note)
            else:
                # 폴백: 본문 끝에 (저자, 연도) + 미주 목록은 참고문헌으로 대체
                last_p.add_run(f" ({s.get('authors','')}, {s.get('year','')})").font.size = Pt(11)

    # ---- 참고문헌 ----
    doc.add_page_break()
    rh = doc.add_heading("참고문헌", level=1)
    rh.text = "참고문헌"
    for ref in citation.reference_list(project["sources"]):
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing = 1.6
        p.paragraph_format.left_indent = Mm(10)
        p.paragraph_format.first_line_indent = Mm(-10)  # 내어쓰기

    # ---- 그림(부록) ----
    import os
    notes = [n for n in project.get("data_notes", []) if os.path.isfile(n.get("figure_file", ""))]
    if notes:
        doc.add_page_break()
        ah = doc.add_heading("부록: 그림", level=1)
        ah.text = "부록: 그림"
        for n in notes:
            try:
                doc.add_picture(n["figure_file"], width=Mm(140))
                cap = doc.add_paragraph(n.get("caption", ""))
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                continue

    if mgr is not None:
        mgr.finalize()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def filename(project) -> str:
    meta = project["meta"]
    dept = meta.get("dept", "국제통상학과")
    sid = meta.get("student_id", "학번")
    name = meta.get("author", "이름")
    return f"({dept}){sid}_{name}.docx"


def render(project, cfg):
    import streamlit as st
    from core import project as P

    st.header("⑦ 내보내기 (.docx)")
    st.caption("학과 양식 자동 적용 · 본인 작성분만 포함 · 각주 인용 + APA 참고문헌")

    st.subheader("표지 정보")
    meta = project["meta"]
    c1, c2 = st.columns(2)
    with c1:
        meta["title"] = st.text_input("제목", value=meta.get("title", ""))
        meta["author"] = st.text_input("이름", value=meta.get("author", ""))
        meta["student_id"] = st.text_input("학번", value=meta.get("student_id", ""))
    with c2:
        meta["advisor"] = st.text_input("지도교수", value=meta.get("advisor", ""))
        meta["dept"] = st.text_input("학과", value=meta.get("dept", "국제통상학과"))
        meta["date"] = st.text_input("제출일", value=meta.get("date", ""))
    if st.button("표지 정보 저장"):
        P.save(project)
        st.success("저장됨")

    st.divider()
    bad = blocked_sections(project)
    if bad:
        st.error("🚫 복붙 차단 절이 있어 내보낼 수 없습니다. M5에서 재작성하세요: " + ", ".join(bad))
        return

    pages = P.estimate_pages(project)
    if pages < 20:
        st.warning(f"본문 추정 {pages}쪽 — 학과 규정(20쪽 이상)에 미달합니다. 내보내기는 가능하나 분량을 보강하세요.")

    if st.button("📄 .docx 생성", type="primary"):
        try:
            data = build_docx(project, use_footnotes=True)
            note = "각주 포함"
        except Exception as e:
            st.warning(f"각주 삽입 실패 → 인용을 본문 표기로 폴백합니다. ({e})")
            data = build_docx(project, use_footnotes=False)
            note = "각주 폴백(본문표기)"
        st.success(f"생성 완료 ({note}). 아래에서 내려받으세요.")
        st.download_button(
            "⬇️ 다운로드", data=data, file_name=filename(project),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    st.info(
        "제출 전 체크: ① 카피킬러로 유사도 30% 이하 확인 → 결과보고서 발급 "
        "② 논문 1부 + 카피킬러 결과보고서를 itrade@dongguk.edu 로 제출 "
        f"③ 파일명: {filename(project)}"
    )
